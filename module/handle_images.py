from docx import Document
from docx.shared import Cm, Inches
from docx2python import docx2python
import os
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image
from io import BytesIO

def get_image_size(image_part):
    """
    Function to extract the size of the image.
    """
    image_bytes = image_part.blob
    image = Image.open(BytesIO(image_bytes))
    
    # Get dimensions in pixels
    width_px, height_px = image.size
    
    # Convert pixels to inches (assuming 96 DPI)
    width_in_inches = width_px / 96
    height_in_inches = height_px / 96
    
    return width_in_inches, height_in_inches

def extract_header_image_sizes(docx_path):
    """
    Extracts image sizes from the header of the DOCX file.
    """
    document = Document(docx_path)
    header = document.sections[0].header  # Assuming single section document

    image_sizes = []

    # Find all image relationships in the header
    for rel in header.part.rels.values():
        if rel.reltype == RT.IMAGE:
            image_part = rel.target_part
            width, height = get_image_size(image_part)
            image_sizes.append((width, height))

    return image_sizes

# Check if the file exists before attempting to remove it
def delete_paragraph(paragraph):
                p = paragraph._element
                p.getparent().remove(p)
                p._p = p._element = None

def extract_images(source_file_path):
        with docx2python(source_file_path) as docx_content:
                docx_content.save_images('./img')

def add_images(trans_file_path, result_file_path, source_file_path):
        
        # trans_file_path = 'other_test_img_result.docx'
        doc = Document(trans_file_path)
        section = doc.sections[0]
        header_para = section.header.paragraphs
        for i in range(len(header_para)):
                if '----media/' in header_para[i].text:                        
                        start_position = header_para[i].text.index('media/')+6
                        end_position = len(header_para[i].text)-5
                        img_file_name = header_para[i].text[start_position:end_position]
                        img_file_path = './img/' + img_file_name
                        print(img_file_path)

                        pp = header_para[i].insert_paragraph_before('')
                        image_sizes = extract_header_image_sizes(source_file_path)
                        for (img_width, img_height) in image_sizes:
                            pp.add_run().add_picture(img_file_path, width=Inches(img_width))    
                       
                        delete_paragraph(header_para[i])
                        if os.path.exists(img_file_path):
                        # Remove the file
                                os.remove(img_file_path)
                                print("File removed successfully.")
                        else:
                                print("The file does not exist.")
        par = doc.paragraphs
        for i in range(len(par)):
                if '----media/' in par[i].text:                        
                        start_position = par[i].text.index('media/')+6
                        end_position = len(par[i].text)-5
                        img_file_name = par[i].text[start_position:end_position]
                        img_file_path = './img/' + img_file_name
                        print(img_file_path)
                        
                        doc1 = Document(source_file_path)
                        for s in doc1.inline_shapes:
                                # print(s._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name)
                                if s._inline.graphic.graphicData.pic.nvPicPr.cNvPr.name == img_file_name:
                                        img_width = round(s.width.cm, 2)
                                        img_height = round(s.height.cm, 2)

                        pp = doc.paragraphs[i].insert_paragraph_before('\n')
                        pp.add_run().add_picture(img_file_path, width=Cm(img_width))
                        # pp.add_run().add_picture(img_file_path, width=Cm(15.92))
                        delete_paragraph(doc.paragraphs[i+1])
                        if os.path.exists(img_file_path):
                        # Remove the file
                                os.remove(img_file_path)
                                print("File removed successfully.")
                        else:
                                print("The file does not exist.")

        # result_file_path = 'demo_better.docx'      

        doc.save(result_file_path)

# add_images(source_file_path='dossier_art_my_studies.docx', trans_file_path='test_img_result.docx', result_file_path='demo_better.docx')