from llama_cpp import Llama
from docx import Document
from docx2python import docx2python
import time
import re
# import handle_footnote
import find_paragraph_footnote
import add_footnote_string


llm = Llama(
    model_path="models/Meta-Llama-3-8B-Instruct-Q8_0.gguf",
    n_ctx=4096,      # Max tokens for in + out
    n_threads=4,     # CPU cores used
    n_gpu_layers=-1,  # Load all layers into VRAM of the GPU
    
)

file_path = 'test.docx'
# file_path = 'test1.docx'

# Load the document
content_doc = Document(file_path)
footnote_doc = docx2python(file_path, html=True)   

trans_lang = "English"
source_lang = "French"

system_prompt = "You are a helpful, smart, kind, and efficient AI assistant. You always fulfill the user's requests to the best of your ability."


def extract_content(doc):
    data_list=[]
    for paragraph in doc.paragraphs:    
        data_list.append(paragraph.text.strip())
    return data_list

def extract_footnote(doc):
    footnote_list = []
    for footnote in doc.footnotes_runs[0][0]:
        for specific in footnote:
            for line in specific:
                split_lines = line.split("\t")
                if "footnote" in split_lines[0]:
                    # footnote_list.append(re.sub(r'\D', '', split_lines[0]))
                    print(re.sub(r'\D', '', split_lines[0]))
                else:
                    footnote_list.append(split_lines[0].strip())
    return footnote_list

def trans_with_ai(prompt, max_tokens=2048):
    """
    Function to send a prompt to the AI and return its response.
    """
    # This function sends the prompt to your AI model and fetches the response
    response = llm(prompt, max_tokens=max_tokens, temperature = 0.2, stop=["Q:", "\n"], echo=False)
    return response

def build_trans_prompt(source_lang, trans_lang, data):    
    prompts = [f"Translate the sentence from {source_lang} to {trans_lang}. I need only translation sentence. Please don't mention about others. Remove unnecessary characters.'{sentence}'" for sentence in data if sentence.strip()]
    return prompts

def build_search_footnote_prompt(paragraph, footnote):    
    prompt = f"Which string of this paragraph can be connected with this footnote text. Please tell me the last 2 to 5 words of that string of this paragraph. The string should be the part of the paragraph not of the footnote text. I need only the string.\nThis is paragraph:'{paragraph}'\nThis is footnote text:'{footnote}'\n "
    return prompt


def main():    
    
    content_data = extract_content(content_doc) 
    content_prompts = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, data=content_data)
    
    footnote_data = extract_footnote(footnote_doc)
    footnote_prompts = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, data=footnote_data)
    
    start_time = time.time()

    result_content_data = []
    result_footnote_data = []

    # translate content of file 
    for prompt in content_prompts:  
        # prompt = system_prompt + prompt        
        if prompt:            
            prompt = f"Q: {prompt}? A: "
            ai_response = trans_with_ai(prompt)           
            result_content_data.append(ai_response['choices'][0]['text'].strip())
        else:
            ai_response = "\n"
            result_content_data.append(ai_response)        
        

    # translate footnote of file
    for prompt in footnote_prompts:
        # prompt = system_prompt + prompt
        prompt = f"Q: {prompt}? A: "
        ai_response = trans_with_ai(prompt)
        print(ai_response)
        result_footnote_data.append(ai_response['choices'][0]['text'].strip())

    end_time = time.time()
    elapsed_time = end_time - start_time
    print(elapsed_time)

    content_tokens_generated = sum(len(sentence.split()) for sentence in result_content_data)
    footnote__tokens_generated = sum(len(sentence.split()) for sentence in result_footnote_data)    


    # Create a new document to store the translated data
    translated_doc = Document()

    # Add translated content sentences to the new document
    for translated_sentence in result_content_data:
        translated_doc.add_paragraph(translated_sentence)   

    # Save the translated document
    content_translated_file_path = 'translated_test_llama3_8b.docx'
    translated_doc.save(content_translated_file_path)   

    footnote_para_indexes = find_paragraph_footnote.paragraphs_for_footnote(file_path=file_path)
    print(footnote_para_indexes)
    i = 0
    refer_string_tokens_generated = 0
    for index in footnote_para_indexes:
        paragraph_text = result_content_data[index]
        print(paragraph_text)
        footnote_text = result_footnote_data[i]
        print(footnote_text)        
        prompt = build_search_footnote_prompt(paragraph=paragraph_text, footnote=footnote_text)        
        # prompt = system_prompt + prompt 
        print(prompt)
        prompt = f"Q: {prompt}? A: "
        # ai_response = trans_with_ai(prompt)
        # print(ai_response)
        refer_string = "!@#$"
        while refer_string not in paragraph_text:
            ai_response = trans_with_ai(prompt)
            refer_string = ai_response['choices'][0]['text'].strip()            
            refer_string = refer_string.replace("'", "")
            print(refer_string)

        refer_string_tokens_generated += len(refer_string.split())
        add_footnote_string.add_footnote(file_path=content_translated_file_path, para_index=index, refer_string=refer_string, footnote_text=footnote_text)
        i += 1  

    end_time = time.time()
    elapsed_time = end_time - start_time
    print(elapsed_time)

    
    total_tokens_generated = content_tokens_generated + footnote__tokens_generated + refer_string_tokens_generated
    # Calculate speed per second
    speed_per_second = total_tokens_generated / elapsed_time
    print("Tokens generated per second:", speed_per_second)

    # Remove the first unnecessary paragraph
    doc = Document(content_translated_file_path)
    def delete_paragraph(paragraph):
        p = paragraph._element
        p.getparent().remove(p)
        p._p = p._element = None
    delete_paragraph(doc.paragraphs[0])
    doc.save(content_translated_file_path)
# To start the chat, call the main_chat function
main()