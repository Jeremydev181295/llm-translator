from llama_cpp import Llama
from docx import Document
import time

llm = Llama(
    model_path="models/Meta-Llama-3-8B-Instruct-Q8_0.gguf",
    n_ctx=4096,      # Max tokens for in + out
    n_threads=4,     # CPU cores used
    n_gpu_layers=-1,  # Load all layers into VRAM of the GPU
    
)

file_path = 'test.docx'
# file_path = 'test1.docx'

# Load the document
doc = Document(file_path)


data=[]
for paragraph in doc.paragraphs:    
    data.append(paragraph.text)
   

trans_lang = "English"
source_lang = "French"

system_prompt = "You are a helpful, smart, kind, and efficient AI assistant. You always fulfill the user's requests to the best of your ability."
prompts = [f"Translate the sentence from {source_lang} to {trans_lang}. I need only translation sentence. Please don't mention about others. '{sentence}'" for sentence in data if sentence.strip()] 
# prompts = [f"Translate the sentence from {source_lang} to {trans_lang}. '{sentence}'" for sentence in data if sentence.strip()]
# prompts = [f"Translate the sentence from {source_lang} to {trans_lang}. '{sentence}'" for sentence in data]

# print(prompts)

def trans_with_ai(prompt, max_tokens=1024):
    """
    Function to send a prompt to the AI and return its response.
    """
    # This function sends the prompt to your AI model and fetches the response
    response = llm(prompt, max_tokens=max_tokens, temperature = 0.1, stop=["Q:", "\n"], echo=False)
    return response

def main_chat():
    start_time = time.time()
    result_data = []

    for prompt in prompts:
        # prompt = system_prompt + prompt           
        prompt = f"Q: {prompt}? A: "
        ai_response = trans_with_ai(prompt)
        print(ai_response)
        result_data.append(ai_response['choices'][0]['text'].strip())

    end_time = time.time()
    elapsed_time = end_time - start_time
    print(elapsed_time)

    total_tokens_generated = sum(len(sentence.split()) for sentence in result_data)

    # Calculate speed per second
    speed_per_second = total_tokens_generated / elapsed_time

    print("Tokens generated per second:", speed_per_second)
    # Create a new document to store the translated data
    translated_doc = Document()

    # Add translated sentences to the new document
    for translated_sentence in result_data:
        translated_doc.add_paragraph(translated_sentence)

    # Save the translated document
    translated_file_path = 'translated_test_llama3_8b.docx'
    translated_doc.save(translated_file_path)

# To start the chat, call the main_chat function
main_chat()