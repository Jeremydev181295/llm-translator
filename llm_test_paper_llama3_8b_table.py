from docx import Document
import llama_model

def contains_table(docx_file):
    doc = Document(docx_file)
    for table in doc.tables:
        return True
    return False

def copy_table(source_path, copy_path):
    source_document = Document(source_path) 
    new_document = Document()
    # Iterate through each table in the source document
    for table in source_document.tables:
        # Add a new paragraph to the new document
        paragraph = new_document.add_paragraph()
        # Add the table to the new paragraph
        paragraph._p.addnext(table._tbl)

    # Save the new document
    new_document.save(copy_path)

def replace_all_text_with_trans_in_table(doc_path, output_path):
    # Load the DOCX file
    doc = Document(doc_path)
    
    # Iterate through all tables in the document
    for table in doc.tables:
        # Iterate through each row in the table
        for row in table.rows:
            # Iterate through each cell in the row
            print(len(row.cells))
            for cell in row.cells:
                if cell.text:
                    prompt = llama_model.build_prompt(cell.text)
                    print(prompt)
                    cell.text = llama_model.response_with_ai(prompt=prompt, temperature=0.001)['choices'][0]['text'].strip()
                    print(cell.text)
    # Save the modified document
    doc.save(output_path)

# Open the source document
def main():
    source_path = 'test.docx'
    copy_path = 'test_output.docx'
    output_path = 'test_output.docx'
    copy_table(source_path, copy_path)
    replace_all_text_with_trans_in_table(copy_path, output_path)


main()


# Example usage



