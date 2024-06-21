from llama_cpp import Llama
from docx import Document
from docx2python import docx2python
import time
from module import handle_footnote
from module import handle_docx
from module import handle_images
from module import handle_format


llm = Llama(
    model_path="models/Meta-Llama-3-8B-Instruct-Q8_0.gguf",
    n_ctx=4096,      # Max tokens for in + out
    n_threads=4,     # CPU cores used
    n_gpu_layers=-1,  # Load all layers into VRAM of the GPU
    
)


def trans_with_ai(prompt, temperature, max_tokens=2048):

    response = llm(prompt, max_tokens=max_tokens, temperature = temperature, stop=["Q:", "\n"], echo=False)
    return response

def build_trans_prompt(source_lang, trans_lang, data):    
    # prompts = [f"Translate from {source_lang} into {trans_lang}. I need only translation sentence. Remove extra symbols and characters. '{sentence}'" for sentence in data if sentence.strip()]
    prompts = [f"'{sentence}'" for sentence in data if sentence.strip()]
    return prompts

# def build_search_footnote_prompt(paragraph, footnote):    
#     prompt = f"Which phrase of this paragraph should be connected with this footnote text. The phrase should be the part of the paragraph not of the footnote text. I need only the phrase and remove other characters.\nThis is paragraph:'{paragraph}'\nThis is footnote text:'{footnote}'\n "
#     return prompt


def main():  
    source_file_path = 'test1.docx'
    handle_images.extract_images(source_file_path=source_file_path)
    # file_path = 'test1.docx'
    read_doc = docx2python(source_file_path)  

    # set source language and translate language
    trans_lang = "English"
    source_lang = "French"
    style = "written"

    # system_prompt = f"Acts as a smart translator. Translate {source_lang} sentences into {trans_lang} sentences in {style} style. Do not remove heading word. If sentence includes '----footnotes----' then translate it. I need only translation sentence."
    system_prompt = f"Acts as a smart translator. Translate into {trans_lang} in {style} style. Keep all words, symbols and original style. Do not mention others. I only need translation sentence."

    # system_prompt = f"Acts as a smart translator. Translate {source_lang} sentences into {trans_lang} sentences in {style} style. Do not remove heading words and translate them. Do not add any characters."

    start_time = time.time()
    #####################################################################################
    header_data = handle_docx.extract_header(read_doc)
    if header_data:
        header_prompts = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, data=header_data)
    else:
        header_prompts = ""

    # translate content of file
    result_header_data = []
    image_flag_header = False
    if header_prompts:
        for prompt in header_prompts:  
            if '----media/' in prompt:
                image_flag_header = True
                result_header_data.append(prompt)
            else:  
                prompt = system_prompt + prompt          
                prompt = f"Q: {prompt} A: "
                ai_response = trans_with_ai(prompt=prompt, temperature=0.1)
                print(ai_response['choices'][0]['text'].strip())           
                result_header_data.append(ai_response['choices'][0]['text'].strip())
    
    translated_header_doc = Document()
    translated_header_doc.sections[0].header.paragraphs[0].text = result_header_data[0]

    # Add translated content sentences to the new document
    for translated_item in result_header_data[1:]:
        print(translated_item)
        translated_header_doc.sections[0].header.add_paragraph(translated_item)
    
    # Save the content translated document
    header_translated_file_path = 'header_translated_test_llama3_8b.docx'
    translated_header_doc.save(header_translated_file_path)

    if image_flag_header == True:
        header_added_file_path = "header_image_added_test.docx"
        # handle_images.extract_images(source_file_path=source_file_path)
        handle_images.add_images(trans_file_path=header_translated_file_path, result_file_path=header_added_file_path, source_file_path=source_file_path)

    # Load the source and destination documents
    source_doc = Document(source_file_path)
    # print(source_doc.styles.)
    target_doc = Document(header_added_file_path)
    handle_format.copy_paragraph_style(source_doc.sections[0].header, target_doc.sections[0].header)
    header_style_result_path = 'header_styled.docx'
    target_doc.save(header_style_result_path)
    ######################################################################################
    # build content prompts
    content_data = handle_docx.extract_content(read_doc) 
    if content_data:
        content_prompts = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, data=content_data)
    else:
        content_prompts = ""   

    # translate content of file
    result_content_data = []
    image_flag = False
    if content_prompts:
        for prompt in content_prompts:  
            if '----media/' in prompt:
                image_flag = True
                result_content_data.append(prompt)
            else:  
                prompt = system_prompt + prompt          
                prompt = f"Q: {prompt} A: "
                ai_response = trans_with_ai(prompt=prompt, temperature=0.1)
                print(ai_response['choices'][0]['text'].strip())           
                result_content_data.append(ai_response['choices'][0]['text'].strip())
    
    # Create a new document to store the translated data
    translated_doc = Document(header_style_result_path)

    # Add translated content sentences to the new document
    for translated_sentence in result_content_data:
        translated_doc.add_paragraph(translated_sentence)   

    # Save the content translated document
    content_translated_file_path = 'content_translated_test_llama3_8b.docx'
    translated_doc.save(content_translated_file_path)

    #######################################################################################
    # # find paragraphs related footnotes and add footnotes              
    # find paragraphs related footnotes
    footnote_para_indexes = handle_footnote.find_paragraphs_for_footnote(file_path=source_file_path)
    print(footnote_para_indexes)
    footnote__tokens_generated = 0

    if footnote_para_indexes != []:
        # build footnote prompts   
        footnote_data = handle_docx.extract_footnote(read_doc)
        if footnote_data:
            footnote_prompts = build_trans_prompt(source_lang=source_lang, trans_lang=trans_lang, data=footnote_data)
        else:
            footnote_prompts = ""

        # translate footnote of file
        result_footnote_data = []
        if footnote_prompts:
            for prompt in footnote_prompts:
                prompt = system_prompt + prompt
                prompt = f"Q: {prompt} A: "
                ai_response = trans_with_ai(prompt=prompt, temperature=0.1)
                print(ai_response)
                result_footnote_data.append(ai_response['choices'][0]['text'].strip())

        # add foottnotes
        i = 0
        for index in footnote_para_indexes:
            footnote_text = result_footnote_data[i]
            i += 1
            refer_string = '----footnote'+str(i)+'----'
            print(refer_string)
            handle_footnote.add_footnote(file_path=content_translated_file_path, para_index=index, refer_string=refer_string, footnote_text=footnote_text)

        # remove unnecessary string created from using spire.doc package    
        revise_doc = Document(content_translated_file_path)
        if revise_doc.paragraphs[0].text == "Evaluation Warning: The document was created with Spire.Doc for Python.":
                handle_docx.delete_paragraph(revise_doc.paragraphs[0])
        
        revise_doc.save(content_translated_file_path)
        footnote__tokens_generated = sum(len(sentence.split()) for sentence in result_footnote_data)
    
    revise_doc = Document(content_translated_file_path)
    handle_docx.remove_string_from_paragraph(revise_doc, '----footnotes----')     
    revise_doc.save(content_translated_file_path)
    #####################################################################################################

    # extract and add images if exist
    if image_flag == True:
        result_file_path = "image_added_test.docx"
        # handle_images.extract_images(source_file_path=source_file_path)
        handle_images.add_images(trans_file_path=content_translated_file_path, result_file_path=result_file_path, source_file_path=source_file_path)

    end_time = time.time()
    elapsed_time = end_time - start_time
    print(elapsed_time)

    content_tokens_generated = sum(len(sentence.split()) for sentence in result_content_data)
        
    total_tokens_generated = content_tokens_generated + footnote__tokens_generated

    # Calculate speed per second
    speed_per_second = total_tokens_generated / elapsed_time
    print("Tokens generated per second:", speed_per_second)


# To start the chat, call the main_chat function
main()