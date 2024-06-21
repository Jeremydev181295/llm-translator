from langchain_community.llms import LlamaCpp
from langchain.prompts import PromptTemplate
from langchain.chains import LLMChain
from docx import Document
import time

# Load the LlamaCpp language model, adjust GPU usage based on your hardware
llm = LlamaCpp(   
    model_path="models/llama-2-7b-chat.Q5_K_M.gguf",    
    # model_path="models/Meta-Llama-3-8B-Instruct-Q6_K.gguf",  
    # model_path="models/Meta-Llama-3-8B-Instruct-Q8_0.gguf", 
    n_gpu_layers=40,
    # n_batch=512,  # Batch size for model processing
    n_batch=1024,
    verbose=False,  # Enable detailed logging for debugging
    device="cuda"
)


# file_path = 'test.docx'
file_path = 'test1.docx'

# Load the document
doc = Document(file_path)


data=[]
for paragraph in doc.paragraphs:    
    data.append(paragraph.text)
   

trans_lang = "English"
source_lang = "French"

# Define the prompt template with a placeholder for the question
template = """
Question: {text}

Answer:
"""

# Generate translation queries for each sentence
translation_queries = [f"Translate '{sentence}' from {source_lang} to {trans_lang}. Don't mention any other sentences." for sentence in data]

# Create an LLMChain to manage interactions with the prompt and model
prompt = PromptTemplate(template=template, input_variables=["text"])
llm_chain = LLMChain(prompt=prompt, llm=llm)

print("Chatbot initialized, ready to translate...")

start_time = time.time()
result_data = []
for query in translation_queries:
    answer = llm_chain.run(query)
    result_data.append(answer)
    # print(answer, '\n')

end_time = time.time()
elapsed_time = end_time - start_time
print(elapsed_time)
# Create a new document to store the translated data
translated_doc = Document()

# Add translated sentences to the new document
for translated_sentence in result_data:
    translated_doc.add_paragraph(translated_sentence)

# Save the translated document
translated_file_path = 'translated_test.docx'
translated_doc.save(translated_file_path)