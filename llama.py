from docx import Document
doc = Document('test.docx')
    
# Iterate through all tables in the document
print(len(doc.tables))
for table in doc.tables:
    print(len(table.rows))
    print('\n')
    # Iterate through each row in the table
    for row in table.rows:
        # Iterate through each cell in the row
        print(len(row.cells))
#         for cell in row.cells:
#             if cell.text:
#                 print(cell.text)
# # Save the modified document
